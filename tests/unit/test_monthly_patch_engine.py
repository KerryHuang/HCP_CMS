"""MonthlyPatchEngine 單元測試。"""
import json
from pathlib import Path

import pytest

from hcp_cms.core.monthly_patch_engine import MonthlyPatchEngine
from hcp_cms.data.database import DatabaseManager


@pytest.fixture
def conn():
    db = DatabaseManager(":memory:")
    db.initialize()
    yield db.connection
    db.connection.close()


class TestLoadIssues:
    def test_load_from_json_file(self, conn, tmp_path):
        data = [
            {"issue_no": "0015659", "program_code": "PAYA001", "program_name": "薪資計算",
             "issue_type": "BugFix", "region": "TW", "description": "修正錯誤",
             "impact": "影響薪資", "test_direction": "執行薪資計算"},
        ]
        f = tmp_path / "issues.json"
        f.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

        eng = MonthlyPatchEngine(conn)
        patch_id = eng.load_issues(source="manual", month_str="202604", file_path=str(f))
        from hcp_cms.data.repositories import PatchRepository
        issues = PatchRepository(conn).list_issues_by_patch(patch_id)
        assert len(issues) == 1
        assert issues[0].issue_no == "0015659"
        assert issues[0].region == "TW"

    def test_load_from_txt_file(self, conn, tmp_path):
        txt = "0015659\tPAYA001\t薪資計算\tBugFix\tTW\t修正錯誤\t影響薪資\t執行薪資計算\n"
        txt += "0015660\tLEAA001\t請假管理\tEnhancement\t共用\t新增功能\t影響請假\t測試請假\n"
        f = tmp_path / "issues.txt"
        f.write_text(txt, encoding="utf-8")

        eng = MonthlyPatchEngine(conn)
        patch_id = eng.load_issues(source="manual", month_str="202604", file_path=str(f))
        from hcp_cms.data.repositories import PatchRepository
        issues = PatchRepository(conn).list_issues_by_patch(patch_id)
        assert len(issues) == 2
        assert issues[1].issue_no == "0015660"

    def test_creates_patch_record(self, conn, tmp_path):
        f = tmp_path / "issues.json"
        f.write_text("[]", encoding="utf-8")
        eng = MonthlyPatchEngine(conn)
        patch_id = eng.load_issues(source="manual", month_str="202604", file_path=str(f))
        from hcp_cms.data.repositories import PatchRepository
        patch = PatchRepository(conn).get_patch_by_id(patch_id)
        assert patch.type == "monthly"
        assert patch.month_str == "202604"


class TestPrepareTestReports:
    def test_detects_and_converts_simplified(self, conn, tmp_path):
        import docx as python_docx
        doc = python_docx.Document()
        doc.add_paragraph("这是简体中文测试报告")
        report_dir = tmp_path / "11G" / "测试报告"
        report_dir.mkdir(parents=True)
        f = report_dir / "01.IP_20260203_0015659_TESTREPORT_11G.docx"
        doc.save(str(f))

        eng = MonthlyPatchEngine(conn)
        result = eng.prepare_test_reports(str(tmp_path))
        assert result["converted"] >= 1

    def test_validates_naming_format(self, conn, tmp_path):
        import docx as python_docx
        doc = python_docx.Document()
        report_dir = tmp_path / "11G" / "測試報告"
        report_dir.mkdir(parents=True)
        bad_name = report_dir / "wrong_name.docx"
        doc.save(str(bad_name))

        eng = MonthlyPatchEngine(conn)
        result = eng.prepare_test_reports(str(tmp_path))
        assert len(result["invalid_names"]) >= 1


class TestGeneratePatchList:
    @pytest.fixture
    def engine_with_patch(self, conn, tmp_path):
        import json
        data = [
            {"issue_no": "0015659", "program_code": "PAYA001", "program_name": "薪資計算",
             "issue_type": "BugFix", "region": "TW", "description": "修正錯誤",
             "impact": "影響薪資", "test_direction": "執行薪資計算"},
            {"issue_no": "0015660", "program_code": "LEAA001", "program_name": "請假管理",
             "issue_type": "Enhancement", "region": "CN", "description": "新增功能",
             "impact": "影響請假", "test_direction": "測試請假"},
        ]
        f = tmp_path / "issues.json"
        f.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        eng = MonthlyPatchEngine(conn)
        pid = eng.load_issues(source="manual", month_str="202604", file_path=str(f))
        return eng, pid

    def test_generates_two_excel_files(self, engine_with_patch, tmp_path):
        eng, pid = engine_with_patch
        paths = eng.generate_patch_list(pid, output_dir=str(tmp_path))
        assert len(paths) == 2
        names = [Path(p).name for p in paths]
        assert any("11G" in n for n in names)
        assert any("12C" in n for n in names)
        for p in paths:
            assert Path(p).exists()

    def test_excel_has_three_tabs(self, engine_with_patch, tmp_path):
        import openpyxl
        eng, pid = engine_with_patch
        paths = eng.generate_patch_list(pid, output_dir=str(tmp_path), month_str="202604")
        wb = openpyxl.load_workbook(paths[0])
        sheet_names = wb.sheetnames
        assert "IT 發行通知" in sheet_names
        assert "HR 發行通知" in sheet_names
        assert "問題修正補充說明" in sheet_names

    def test_region_color_coding(self, engine_with_patch, tmp_path):
        import openpyxl
        eng, pid = engine_with_patch
        paths = eng.generate_patch_list(pid, output_dir=str(tmp_path), month_str="202604")
        wb = openpyxl.load_workbook(paths[0])
        ws = wb["HR 發行通知"]
        tw_fill = ws.cell(2, 2).fill.fgColor.rgb  # 第一筆 TW
        cn_fill = ws.cell(3, 2).fill.fgColor.rgb  # 第二筆 CN
        assert tw_fill != cn_fill


class TestGenerateNotifyHtml:
    @pytest.fixture
    def engine_with_patch(self, conn, tmp_path):
        import json
        data = [{"issue_no": "0015659", "program_code": "PAYA001", "program_name": "薪資計算",
                 "issue_type": "BugFix", "region": "TW", "description": "修正錯誤",
                 "impact": "影響薪資", "test_direction": "執行薪資計算"}]
        f = tmp_path / "issues.json"
        f.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        eng = MonthlyPatchEngine(conn)
        pid = eng.load_issues(source="manual", month_str="202604", file_path=str(f))
        return eng, pid

    def test_generates_html_file(self, engine_with_patch, tmp_path):
        eng, pid = engine_with_patch
        path = eng.generate_notify_html(pid, output_dir=str(tmp_path), month_str="202604")
        assert Path(path).exists()
        assert path.endswith(".html")

    def test_html_contains_issue_no(self, engine_with_patch, tmp_path):
        eng, pid = engine_with_patch
        path = eng.generate_notify_html(pid, output_dir=str(tmp_path), month_str="202604")
        content = Path(path).read_text(encoding="utf-8")
        assert "0015659" in content

    def test_html_filename_format(self, engine_with_patch, tmp_path):
        eng, pid = engine_with_patch
        path = eng.generate_notify_html(pid, output_dir=str(tmp_path), month_str="202604")
        assert "202604" in Path(path).name
        assert "大PATCH更新通知" in Path(path).name


class TestScanMonthlyDir:
    def test_engine_stores_conn(self, conn):
        eng = MonthlyPatchEngine(conn)
        assert eng._conn is conn

    def test_detect_structure_mode_a(self, conn, tmp_path):
        (tmp_path / "11G").mkdir()
        eng = MonthlyPatchEngine(conn)
        assert eng._detect_structure(tmp_path) == "A"

    def test_detect_structure_mode_b(self, conn, tmp_path):
        (tmp_path / "01.IP_20241128_0016552_11G.7z").write_bytes(b"")
        eng = MonthlyPatchEngine(conn)
        assert eng._detect_structure(tmp_path) == "B"

    def test_reorganize_to_mode_a(self, conn, tmp_path):
        (tmp_path / "01.IP_20241128_0016552_11G.7z").write_bytes(b"dummy")
        (tmp_path / "01.IP_20241128_0016552_12C.zip").write_bytes(b"dummy")
        (tmp_path / "01.IP_20241204_0016552_TESTREPORT_11G.doc").write_bytes(b"dummy")
        (tmp_path / "01.IP_20241204_0016552_TESTREPORT_12C.doc").write_bytes(b"dummy")

        eng = MonthlyPatchEngine(conn)
        eng._reorganize_to_mode_a(tmp_path)

        assert (tmp_path / "11G" / "01.IP_20241128_0016552_11G.7z").exists()
        assert (tmp_path / "12C" / "01.IP_20241128_0016552_12C.zip").exists()
        assert (tmp_path / "11G" / "測試報告" / "01.IP_20241204_0016552_TESTREPORT_11G.doc").exists()
        assert (tmp_path / "12C" / "測試報告" / "01.IP_20241204_0016552_TESTREPORT_12C.doc").exists()

    def test_scan_monthly_dir_mode_a(self, conn, tmp_path, monkeypatch):
        import json
        import py7zr
        from hcp_cms.core.monthly_patch_engine import MonthlyPatchEngine
        from hcp_cms.data.repositories import PatchRepository

        # 建立模式 A 結構：11G/ 含一個 .7z
        ver_dir = tmp_path / "11G"
        ver_dir.mkdir()
        archive = ver_dir / "01.IP_20241128_0016552_11G.7z"
        # 建立包含 form/ sql/ 的 .7z
        extracted_stub = tmp_path / "_stub"
        extracted_stub.mkdir()
        (extracted_stub / "form").mkdir()
        (extracted_stub / "form" / "HRWF304.fmb").write_bytes(b"")
        (extracted_stub / "sql").mkdir()
        (extracted_stub / "sql" / "pk_test.sql").write_bytes(b"")
        with py7zr.SevenZipFile(str(archive), "w") as z:
            for f in extracted_stub.rglob("*"):
                if f.is_file():
                    z.write(str(f), str(f.relative_to(extracted_stub)))

        # mock _read_release_note 回傳固定資料
        def fake_release(self_inner, path):
            return [{"issue_no": "0016552", "issue_type": "BugFix",
                     "description": "測試說明", "region": "共用"}]
        monkeypatch.setattr(MonthlyPatchEngine, "_read_release_note", fake_release)

        eng = MonthlyPatchEngine(conn)
        result = eng.scan_monthly_dir(str(tmp_path), "202412")

        assert "11G" in result
        repo = PatchRepository(conn)
        issues = repo.list_issues_by_patch(result["11G"])
        assert len(issues) == 1
        assert issues[0].issue_no == "0016552"
        meta = json.loads(issues[0].mantis_detail)
        assert "HRWF304" in meta["form_files"]
        assert "pk_test" in meta["sql_files"]

    def test_scan_monthly_dir_mode_b_reorganizes(self, conn, tmp_path, monkeypatch):
        from hcp_cms.core.monthly_patch_engine import MonthlyPatchEngine

        archive_11g = tmp_path / "01.IP_20241128_0016552_11G.7z"
        archive_11g.write_bytes(b"dummy")
        report = tmp_path / "01.IP_20241204_0016552_TESTREPORT_11G.doc"
        report.write_bytes(b"dummy")

        def fake_extract(self_inner, archive, extract_dir):
            extract_dir.mkdir(parents=True, exist_ok=True)

        def fake_release(self_inner, path):
            return []

        monkeypatch.setattr(MonthlyPatchEngine, "_extract_archive", fake_extract)
        monkeypatch.setattr(MonthlyPatchEngine, "_read_release_note", fake_release)

        eng = MonthlyPatchEngine(conn)
        eng.scan_monthly_dir(str(tmp_path), "202412")

        assert (tmp_path / "11G" / "01.IP_20241128_0016552_11G.7z").exists()
        assert (tmp_path / "11G" / "測試報告" / "01.IP_20241204_0016552_TESTREPORT_11G.doc").exists()


class TestGeneratePatchListFromDir:
    @pytest.fixture
    def conn(self):
        db = DatabaseManager(":memory:")
        db.initialize()
        yield db.connection
        db.connection.close()

    def test_find_test_report_found(self, conn, tmp_path):
        (tmp_path / "11G" / "測試報告").mkdir(parents=True)
        report = tmp_path / "11G" / "測試報告" / "01.IP_20241204_0016552_TESTREPORT_11G.doc"
        report.write_bytes(b"")

        eng = MonthlyPatchEngine(conn)
        result = eng._find_test_report(tmp_path, "11G", "0016552")
        assert result is not None
        assert "0016552" in result

    def test_find_test_report_not_found(self, conn, tmp_path):
        (tmp_path / "11G" / "測試報告").mkdir(parents=True)
        eng = MonthlyPatchEngine(conn)
        assert eng._find_test_report(tmp_path, "11G", "9999999") is None

    def test_generate_patch_list_from_dir_creates_files(self, conn, tmp_path):
        import json
        from hcp_cms.data.repositories import PatchRepository

        repo = PatchRepository(conn)
        from hcp_cms.data.models import PatchRecord, PatchIssue
        pid_11g = repo.insert_patch(PatchRecord(type="monthly", month_str="202412", patch_dir=str(tmp_path)))
        repo.insert_issue(PatchIssue(
            patch_id=pid_11g, issue_no="0016552", issue_type="BugFix",
            region="共用", description="測試問題修正", source="scan",
            program_code="HRWF304", program_name="派退宿功能",
            mantis_detail=json.dumps({
                "form_files": ["HRWF304"], "sql_files": [], "muti_files": [],
                "archive_name": "01.IP_20241128_0016552_11G.7z"
            })
        ))
        pid_12c = repo.insert_patch(PatchRecord(type="monthly", month_str="202412", patch_dir=str(tmp_path)))
        repo.insert_issue(PatchIssue(
            patch_id=pid_12c, issue_no="0016552", issue_type="BugFix",
            region="共用", description="測試問題修正", source="scan",
            mantis_detail=json.dumps({"form_files": [], "sql_files": [], "muti_files": []})
        ))

        (tmp_path / "11G").mkdir()
        (tmp_path / "12C").mkdir()
        (tmp_path / "11G" / "測試報告").mkdir()

        eng = MonthlyPatchEngine(conn)
        paths = eng.generate_patch_list_from_dir(
            {"11G": pid_11g, "12C": pid_12c}, str(tmp_path), "202412"
        )

        assert len(paths) == 2
        assert any("11G" in p for p in paths)
        assert any("12C" in p for p in paths)
        for p in paths:
            assert Path(p).exists()

    def test_generate_sheet_names(self, conn, tmp_path):
        import json
        from openpyxl import load_workbook
        from hcp_cms.data.repositories import PatchRepository
        from hcp_cms.data.models import PatchRecord, PatchIssue

        repo = PatchRepository(conn)
        pid = repo.insert_patch(PatchRecord(type="monthly", month_str="202412", patch_dir=str(tmp_path)))
        repo.insert_issue(PatchIssue(
            patch_id=pid, issue_no="0016552", source="scan",
            mantis_detail=json.dumps({"form_files": [], "sql_files": [], "muti_files": []})
        ))
        (tmp_path / "11G").mkdir()

        eng = MonthlyPatchEngine(conn)
        paths = eng.generate_patch_list_from_dir({"11G": pid}, str(tmp_path), "202412")

        wb = load_workbook(paths[0])
        assert "IT 發行通知" in wb.sheetnames
        assert "HR 發行通知" in wb.sheetnames
        assert "問題修正補充說明" in wb.sheetnames

    def test_generate_it_sheet_form_files(self, conn, tmp_path):
        import json
        from openpyxl import load_workbook
        from hcp_cms.data.repositories import PatchRepository
        from hcp_cms.data.models import PatchRecord, PatchIssue

        repo = PatchRepository(conn)
        pid = repo.insert_patch(PatchRecord(type="monthly", month_str="202412", patch_dir=str(tmp_path)))
        repo.insert_issue(PatchIssue(
            patch_id=pid, issue_no="0016552", issue_type="BugFix",
            region="共用", description="修正說明", source="scan",
            mantis_detail=json.dumps({"form_files": ["HRWF304"], "sql_files": ["pk_test"], "muti_files": []})
        ))
        (tmp_path / "11G").mkdir()

        eng = MonthlyPatchEngine(conn)
        paths = eng.generate_patch_list_from_dir({"11G": pid}, str(tmp_path), "202412")

        wb = load_workbook(paths[0])
        ws = wb["IT 發行通知"]
        # Row 2, col 5 = FORM目錄
        assert ws.cell(2, 5).value == "HRWF304"
        # Row 2, col 6 = DB物件
        assert ws.cell(2, 6).value == "pk_test"

    def test_generate_test_report_hyperlink(self, conn, tmp_path):
        import json
        from openpyxl import load_workbook
        from hcp_cms.data.repositories import PatchRepository
        from hcp_cms.data.models import PatchRecord, PatchIssue

        repo = PatchRepository(conn)
        pid = repo.insert_patch(PatchRecord(type="monthly", month_str="202412", patch_dir=str(tmp_path)))
        repo.insert_issue(PatchIssue(
            patch_id=pid, issue_no="0016552", source="scan",
            mantis_detail=json.dumps({"form_files": [], "sql_files": [], "muti_files": []})
        ))
        (tmp_path / "11G").mkdir()
        report_dir = tmp_path / "11G" / "測試報告"
        report_dir.mkdir()
        (report_dir / "01.IP_20241204_0016552_TESTREPORT_11G.doc").write_bytes(b"")

        eng = MonthlyPatchEngine(conn)
        paths = eng.generate_patch_list_from_dir({"11G": pid}, str(tmp_path), "202412")

        wb = load_workbook(paths[0])
        ws = wb["問題修正補充說明"]
        # Row 2 = 第一筆 issue，col 2 = 測試報告（含超連結）
        cell = ws.cell(2, 2)
        assert cell.hyperlink is not None
        hyperlink_target = cell.hyperlink.target if hasattr(cell.hyperlink, "target") else str(cell.hyperlink)
        assert "0016552" in hyperlink_target


class TestGenerateNotifyHtmlFromDir:
    @pytest.fixture
    def conn(self):
        db = DatabaseManager(":memory:")
        db.initialize()
        yield db.connection
        db.connection.close()

    def test_generates_html_per_version(self, conn, tmp_path):
        from pathlib import Path
        from hcp_cms.data.repositories import PatchRepository
        from hcp_cms.data.models import PatchRecord, PatchIssue

        repo = PatchRepository(conn)
        pid_11g = repo.insert_patch(PatchRecord(type="monthly", month_str="202412", patch_dir=str(tmp_path)))
        repo.insert_issue(PatchIssue(patch_id=pid_11g, issue_no="0016552", source="scan"))
        pid_12c = repo.insert_patch(PatchRecord(type="monthly", month_str="202412", patch_dir=str(tmp_path)))
        repo.insert_issue(PatchIssue(patch_id=pid_12c, issue_no="0016552", source="scan"))

        eng = MonthlyPatchEngine(conn)
        paths = eng.generate_notify_html_from_dir(
            {"11G": pid_11g, "12C": pid_12c}, str(tmp_path), "202412"
        )

        assert len(paths) == 2
        assert any("11G" in p for p in paths)
        assert any("12C" in p for p in paths)
        for p in paths:
            assert Path(p).exists()
