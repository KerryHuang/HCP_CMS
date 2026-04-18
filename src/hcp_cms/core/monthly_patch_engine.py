"""MonthlyPatchEngine — 每月大 PATCH 整理流程。"""

from __future__ import annotations

import json
import logging
import re
import sqlite3
from collections.abc import Callable
from pathlib import Path

from hcp_cms.data.models import PatchIssue, PatchRecord
from hcp_cms.data.repositories import PatchRepository
from hcp_cms.services.claude_content import ClaudeContentService
from hcp_cms.services.credential import CredentialManager
from hcp_cms.services.mantis.soap import MantisSoapClient

_TXT_FIELDS = [
    "issue_no",
    "program_code",
    "program_name",
    "issue_type",
    "region",
    "description",
    "impact",
    "test_direction",
]

_SEASON_COLORS = {
    1: ("1F4E79", "2E75B6", "🌸"),
    2: ("1F4E79", "2E75B6", "🌸"),
    3: ("1F4E79", "2E75B6", "🌸"),
    4: ("1B5E20", "2E7D32", "🌿"),
    5: ("1B5E20", "2E7D32", "🌿"),
    6: ("1B5E20", "2E7D32", "🌿"),
    7: ("E65100", "F57C00", "🌻"),
    8: ("E65100", "F57C00", "🌻"),
    9: ("E65100", "F57C00", "🌻"),
    10: ("263238", "37474F", "❄️"),
    11: ("263238", "37474F", "❄️"),
    12: ("263238", "37474F", "❄️"),
}


class MonthlyPatchEngine:
    def __init__(self, conn: sqlite3.Connection) -> None:
        self._conn = conn
        self._repo = PatchRepository(conn)

    # ── Issue 載入 ───────────────────────────────────────────────────────────

    def load_issues(self, source: str, month_str: str, file_path: str | None = None) -> int:
        """載入 Issue 清單，回傳 patch_id。

        source='manual': 讀 file_path（.json 或 .txt tab 分隔）
        source='mantis': 透過 PlaywrightMantisService（需另行呼叫）
        """
        patch = PatchRecord(type="monthly", month_str=month_str)
        patch_id = self._repo.insert_patch(patch)

        if source == "manual" and file_path:
            issues = self._load_file(file_path)
            for idx, raw in enumerate(issues):
                issue = PatchIssue(
                    patch_id=patch_id,
                    issue_no=raw.get("issue_no", ""),
                    program_code=raw.get("program_code"),
                    program_name=raw.get("program_name"),
                    issue_type=raw.get("issue_type", "BugFix"),
                    region=raw.get("region", "共用"),
                    description=raw.get("description"),
                    impact=raw.get("impact"),
                    test_direction=raw.get("test_direction"),
                    source="manual",
                    sort_order=idx,
                )
                self._repo.insert_issue(issue)

        return patch_id

    def _load_file(self, file_path: str) -> list[dict]:
        path = Path(file_path)
        if path.suffix.lower() == ".json":
            return json.loads(path.read_text(encoding="utf-8"))
        rows = []
        for line in path.read_text(encoding="utf-8").splitlines():
            parts = line.rstrip("\n").split("\t")
            if len(parts) >= len(_TXT_FIELDS):
                rows.append(dict(zip(_TXT_FIELDS, parts)))
        return rows

    def get_issue_count(self, patch_id: int) -> int:
        """回傳 Patch 的 Issue 筆數。"""
        return len(self._repo.list_issues_by_patch(patch_id))

    def get_issues(self, patch_id: int) -> list[PatchIssue]:
        """回傳 Patch 的 Issue 清單。"""
        return self._repo.list_issues_by_patch(patch_id)

    # ── 資料夾掃描 ────────────────────────────────────────────────────────────

    _ARCHIVE_RE = re.compile(r"\d{2}\.IP_\d{8}_(\d{7})_", re.IGNORECASE)

    def _detect_structure(self, base: Path) -> str:
        """回傳 'A'（有 11G/12C 子目錄）或 'B'（平坦）。"""
        if (base / "11G").exists() or (base / "12C").exists():
            return "A"
        return "B"

    def _reorganize_to_mode_a(self, base: Path) -> None:
        """將模式 B 平坦結構整理至 11G/12C 子目錄。"""
        import shutil
        for version in ("11G", "12C"):
            ver_dir = base / version
            report_dir = ver_dir / "測試報告"
            ver_dir.mkdir(exist_ok=True)
            report_dir.mkdir(exist_ok=True)
            v_upper = version.upper()
            for f in list(base.iterdir()):
                if not f.is_file():
                    continue
                name_upper = f.name.upper()
                if f"_{v_upper}" not in name_upper:
                    continue
                if "TESTREPORT" in name_upper:
                    shutil.move(str(f), str(report_dir / f.name))
                elif f.suffix.lower() in (".7z", ".zip", ".rar"):
                    shutil.move(str(f), str(ver_dir / f.name))

    def _extract_archive(self, archive: Path, extract_dir: Path) -> None:
        """解壓 .7z 或 .zip 至 extract_dir。"""
        extract_dir.mkdir(parents=True, exist_ok=True)
        suffix = archive.suffix.lower()
        if suffix == ".7z":
            import py7zr
            with py7zr.SevenZipFile(str(archive), mode="r") as z:
                z.extractall(path=str(extract_dir))
        elif suffix == ".zip":
            import zipfile
            with zipfile.ZipFile(str(archive), "r") as z:
                z.extractall(path=str(extract_dir))
        else:
            logging.warning("不支援的封存格式，略過 [%s]", archive.name)

    def _list_files(self, directory: Path, extensions: list[str], keep_ext: bool = False) -> list[str]:
        """列出目錄內符合副檔名的檔案，keep_ext=False 則去副檔名。"""
        if not directory.exists():
            return []
        return [
            f.name if keep_ext else f.stem
            for f in sorted(directory.iterdir())
            if f.is_file() and f.suffix.lower() in extensions
        ]

    def _extract_issue_no(self, filename: str) -> str | None:
        """從封存檔名解析 7 位數 Issue No，如 '01.IP_20241128_0016552_11G.7z' → '0016552'。"""
        m = self._ARCHIVE_RE.search(filename)
        return m.group(1) if m else None

    def _find_extract_root(self, extract_dir: Path) -> Path:
        """若解壓資料夾只有一個子資料夾（封存內頂層目錄），回傳該子資料夾；否則回傳 extract_dir。"""
        if not extract_dir.exists():
            return extract_dir
        subdirs = [d for d in extract_dir.iterdir() if d.is_dir()]
        if len(subdirs) == 1 and not any(f for f in extract_dir.iterdir() if f.is_file()):
            return subdirs[0]
        return extract_dir

    def _read_release_note(self, extract_dir: Path) -> list[dict[str, str]]:
        """在解壓資料夾尋找 ReleaseNote，委託 SinglePatchEngine 解析。"""
        if not extract_dir.exists():
            return []
        from hcp_cms.core.patch_engine import SinglePatchEngine
        for f in extract_dir.iterdir():
            if not f.is_file():
                continue
            low = f.name.lower()
            if "releasenote" in low or "release_note" in low:
                return SinglePatchEngine(self._conn).read_release_doc(str(f))
        return []

    def scan_monthly_dir(
        self,
        patch_dir: str,
        month_str: str,
        progress: Callable[[str], None] | None = None,
    ) -> dict[str, int]:
        """掃描月份資料夾，建立各版本 PatchRecord 並匯入 Issues。

        自動偵測模式 A（11G/12C 子目錄）或模式 B（平坦），模式 B 時自動整理。
        回傳 {"11G": patch_id, "12C": patch_id}（版本不存在則無對應 key）。
        progress: 進度回調，傳入要顯示的訊息字串。
        """
        def _log(msg: str) -> None:
            if progress:
                progress(msg)

        base = Path(patch_dir)
        patch_ids: dict[str, int] = {}

        if self._detect_structure(base) == "B":
            _log("🔄 偵測到平坦結構，自動整理為 11G/12C 子目錄…")
            self._reorganize_to_mode_a(base)

        for version in ("11G", "12C"):
            version_dir = base / version
            if not version_dir.exists():
                continue
            archives = sorted(version_dir.glob("*.7z")) + sorted(version_dir.glob("*.zip"))
            if not archives:
                _log(f"⚠️ {version}/ 目錄下無 .7z / .zip 封存檔，略過")
                continue

            _log(f"📂 {version}：找到 {len(archives)} 個封存檔")
            patch = PatchRecord(type="monthly", month_str=month_str, patch_dir=patch_dir)
            patch_id = self._repo.insert_patch(patch)
            patch_ids[version] = patch_id

            for idx, archive in enumerate(archives):
                issue_no = self._extract_issue_no(archive.name)
                _log(f"  [{idx + 1}/{len(archives)}] 解壓縮 {archive.name}…")
                extract_dir = version_dir / f"_extracted_{archive.stem}"
                try:
                    self._extract_archive(archive, extract_dir)
                except Exception as e:
                    logging.warning("解壓縮失敗 [%s]: %s", archive.name, e)
                    _log(f"  ❌ 解壓縮失敗：{e}")
                    continue

                root = self._find_extract_root(extract_dir)
                raw_issues = self._read_release_note(root)
                form_files = self._list_files(root / "form", [".fmb", ".rdf", ".fmx"])
                sql_files = self._list_files(root / "sql", [".sql"])
                muti_files = self._list_files(root / "muti", [".sql"], keep_ext=True)
                scan_meta = json.dumps({
                    "form_files": form_files,
                    "sql_files": sql_files,
                    "muti_files": muti_files,
                    "archive_name": archive.name,
                })

                if raw_issues:
                    _log(f"  ✅ {archive.name} → {len(raw_issues)} 筆 Issue（來自 ReleaseNote）")
                    for raw_idx, raw in enumerate(raw_issues):
                        self._repo.insert_issue(PatchIssue(
                            patch_id=patch_id,
                            issue_no=raw.get("issue_no", issue_no or ""),
                            issue_type=raw.get("issue_type", "BugFix"),
                            region=raw.get("region", "共用"),
                            description=raw.get("description"),
                            source="scan",
                            sort_order=idx * 100 + raw_idx,
                            mantis_detail=scan_meta,
                        ))
                elif issue_no:
                    _log(f"  ✅ {archive.name} → Issue {issue_no}（無 ReleaseNote，從檔名解析）")
                    self._repo.insert_issue(PatchIssue(
                        patch_id=patch_id,
                        issue_no=issue_no,
                        source="scan",
                        sort_order=idx,
                        mantis_detail=scan_meta,
                    ))
                else:
                    _log(f"  ⚠️ {archive.name} → 無法解析 Issue No，略過")

        if not patch_ids:
            _log("⚠️ 未找到任何版本（11G/12C）的封存檔，請確認資料夾結構")

        return patch_ids

    # ── 測試報告整理 ─────────────────────────────────────────────────────────

    _REPORT_NAME_RE = re.compile(
        r"^\d{2}\.IP_\d{8}_\d{7}_TESTREPORT_(11G|12C)\.(doc|docx)$",
        re.IGNORECASE,
    )

    def prepare_test_reports(self, month_dir: str) -> dict:
        """掃描測試報告資料夾，轉換簡體→繁體，驗證命名格式。"""
        base = Path(month_dir)
        result: dict = {"converted": 0, "invalid_names": [], "files_checked": 0}

        for version in ("11G", "12C"):
            for sub in ("測試報告", "测试报告"):
                report_dir = base / version / sub
                if not report_dir.exists():
                    continue
                for f in report_dir.iterdir():
                    if f.suffix.lower() not in (".doc", ".docx"):
                        continue
                    result["files_checked"] += 1
                    if not self._REPORT_NAME_RE.match(f.name):
                        result["invalid_names"].append(f.name)
                    if f.suffix.lower() == ".docx":
                        converted = self._convert_simplified_to_traditional(f)
                        if converted:
                            result["converted"] += 1

        return result

    def _convert_simplified_to_traditional(self, path: Path) -> bool:
        """若 .docx 含簡體字則就地轉換為繁體，回傳是否有轉換。"""
        try:
            import docx as python_docx
            import opencc

            cc = opencc.OpenCC("s2t")
            doc = python_docx.Document(str(path))
            changed = False
            for para in doc.paragraphs:
                converted = cc.convert(para.text)
                if converted != para.text:
                    for run in para.runs:
                        run.text = cc.convert(run.text)
                    changed = True
            if changed:
                doc.save(str(path))
            return changed
        except Exception as e:
            logging.warning("opencc 轉換失敗 [%s]: %s", path.name, e)
            return False

    _FETCH_NO_CONN  = -1   # Mantis 連線失敗
    _FETCH_NO_ISSUE = -2   # 無 Issue 可處理

    def fetch_supplements(
        self,
        patch_id: int,
        progress: Callable[[str], None] | None = None,
    ) -> int:
        """從 Mantis 取得各 Issue 說明，以 Claude 整理補充說明五欄位。

        回傳值：
            >= 0  → 成功更新筆數
            -1    → Mantis 連線失敗
            -2    → 該 Patch 無 Issue
        """
        def _log(msg: str) -> None:
            if progress:
                progress(msg)

        client = self._build_mantis_client()
        if client is None:
            return self._FETCH_NO_CONN
        svc = ClaudeContentService()
        issues = self._repo.list_issues_by_patch(patch_id)
        if not issues:
            return self._FETCH_NO_ISSUE
        count = 0
        for iss in issues:
            mantis_id = iss.issue_no.lstrip("0") or "0"
            _log(f"  🔍 查詢 Issue {iss.issue_no} (Mantis id={mantis_id})…")
            supplement = self._fetch_supplement(iss.issue_no, client, svc)
            if not any(supplement.values()):
                _log(f"  ⚠️ Issue {iss.issue_no}：Mantis 無資料（{client.last_error or '補充欄位為空'}）")
                continue
            existing = self._parse_scan_meta(iss)
            existing["supplement"] = supplement
            self._repo.update_issue_mantis_detail(iss.issue_id, json.dumps(existing, ensure_ascii=False))
            _log(f"  ✅ Issue {iss.issue_no}：補充說明已更新")
            count += 1
        return count

    def _fetch_supplement(
        self, issue_no: str, client: MantisSoapClient, svc: ClaudeContentService
    ) -> dict[str, str]:
        """呼叫 Mantis + Claude，回傳補充說明五欄位 dict。"""
        empty = {"修改原因": "", "原問題": "", "範例說明": "", "修正後": "", "注意事項": ""}
        try:
            issue = client.get_issue(issue_no.lstrip("0") or "0")
            if issue is None:
                logging.warning("fetch_supplement: Mantis 找不到 issue_no=%s: %s",
                                issue_no, client.last_error)
                return empty
            notes_text = "\n".join(n.text for n in (issue.notes_list or []))
            full_text = f"{issue.description}\n{notes_text}".strip()
            return svc.extract_supplement(full_text)
        except Exception as e:
            logging.warning("fetch_supplement 失敗 [%s]: %s", issue_no, e)
            return empty

    def _build_mantis_client(self) -> MantisSoapClient | None:
        """從 keyring 讀取憑證建立 MantisSoapClient，失敗時回傳 None。"""
        try:
            from urllib.parse import urlparse
            creds = CredentialManager()
            url = creds.retrieve("mantis_url") or ""
            user = creds.retrieve("mantis_user") or ""
            pwd = creds.retrieve("mantis_password") or ""
            if not url:
                return None
            parsed = urlparse(url)
            path = parsed.path
            if ".php" in path:
                path = path[:path.rfind("/")]
            base_url = f"{parsed.scheme}://{parsed.netloc}{path}".rstrip("/")
            client = MantisSoapClient(base_url, user, pwd)
            if not client.connect():
                return None
            return client
        except Exception:
            return None

    def run_s2t(self, scan_dir: str) -> dict[str, int]:
        """掃描 scan_dir 下所有版本子目錄的測試報告資料夾，將 .docx 簡體轉繁體。
        回傳 {filename: converted_char_count}，0 表示無需轉換。
        """
        import docx as python_docx
        import opencc

        base = Path(scan_dir)
        result: dict[str, int] = {}
        cc = opencc.OpenCC("s2t")

        for report_dir in base.rglob("測試報告"):
            if not report_dir.is_dir():
                continue
            for docx_path in sorted(report_dir.glob("*.docx")):
                try:
                    doc = python_docx.Document(str(docx_path))
                    changed_chars = 0
                    for para in doc.paragraphs:
                        converted = cc.convert(para.text)
                        if converted != para.text:
                            diff = sum(1 for a, b in zip(para.text, converted) if a != b)
                            changed_chars += diff
                            for run in para.runs:
                                run.text = cc.convert(run.text)
                    if changed_chars:
                        doc.save(str(docx_path))
                    result[docx_path.name] = changed_chars
                except Exception as e:
                    logging.warning("S2T 失敗 [%s]: %s", docx_path.name, e)
                    result[docx_path.name] = -1
        return result

    # ── PATCH_LIST Excel ─────────────────────────────────────────────────────

    _CLR_TW = "D6EAF8"  # TW 淡藍
    _CLR_CN = "FFF3CD"  # CN 淡橘黃
    _CLR_BOTH = "E8F5E9"  # 共用 淡綠
    _CLR_BUG = "FCE4D6"  # BugFix
    _CLR_ENH = "E2EFDA"  # Enhancement
    _HDR_DARK = "1F3864"  # 主標題深藍
    _HDR_MID = "2E75B6"  # 副標題中藍

    def generate_patch_list(self, patch_id: int, output_dir: str, month_str: str | None = None) -> list[str]:
        """產 PATCH_LIST_{YYYYMM}_11G.xlsx 與 12C.xlsx，各含 3 頁籤。"""
        from openpyxl import Workbook
        from openpyxl.styles import PatternFill

        issues = self._repo.list_issues_by_patch(patch_id)
        if month_str is None:
            patch = self._repo.get_patch_by_id(patch_id)
            month_str = patch.month_str if patch and patch.month_str else "000000"

        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)
        paths = []

        for version in ("11G", "12C"):
            wb = Workbook()

            # 頁籤①：IT 發行通知（8 欄）
            ws_it = wb.active
            ws_it.title = "IT 發行通知"
            hdrs_it = ["Issue No", "類型", "程式代號", "說明", "FORM目錄", "DB物件", "多語更新", "備註"]
            self._write_patch_header(ws_it, hdrs_it)
            for i, iss in enumerate(issues, start=2):
                ws_it.cell(i, 1).value = iss.issue_no
                ws_it.cell(i, 2).value = iss.issue_type
                ws_it.cell(i, 3).value = iss.program_code
                ws_it.cell(i, 4).value = iss.description
                row_fill = self._CLR_ENH if iss.issue_type == "Enhancement" else self._CLR_BUG
                for c in range(1, 9):
                    ws_it.cell(i, c).fill = PatternFill("solid", fgColor=row_fill)

            # 頁籤②：HR 發行通知（11 欄）
            ws_hr = wb.create_sheet("HR 發行通知")
            hdrs_hr = [
                "Issue No",
                "計區域",
                "類型",
                "程式代號",
                "程式名稱",
                "功能說明",
                "影響說明/用途",
                "相關程式(FORM)",
                "上線所需動作",
                "測試方向及注意事項",
                "備註",
            ]
            self._write_patch_header(ws_hr, hdrs_hr)
            for i, iss in enumerate(issues, start=2):
                region_fill = {"TW": self._CLR_TW, "CN": self._CLR_CN}.get(iss.region, self._CLR_BOTH)
                ws_hr.cell(i, 1).value = iss.issue_no
                ws_hr.cell(i, 2).value = iss.region
                ws_hr.cell(i, 3).value = iss.issue_type
                ws_hr.cell(i, 4).value = iss.program_code
                ws_hr.cell(i, 5).value = iss.program_name
                ws_hr.cell(i, 6).value = iss.description
                ws_hr.cell(i, 7).value = iss.impact
                ws_hr.cell(i, 9).value = "請與資訊單位確認是否已完成更新，確認更新完成再進行測試"
                ws_hr.cell(i, 10).value = iss.test_direction
                for c in range(1, 12):
                    ws_hr.cell(i, c).fill = PatternFill("solid", fgColor=region_fill)

            # 頁籤③：問題修正補充說明
            ws_supp = wb.create_sheet("問題修正補充說明")
            self._write_patch_header(ws_supp, ["Issue No", "修改原因", "原問題", "範例說明", "修正後", "注意事項"])
            for i, iss in enumerate(issues, start=2):
                ws_supp.cell(i, 1).value = iss.issue_no
                if iss.mantis_detail:
                    try:
                        detail = json.loads(iss.mantis_detail)
                        ws_supp.cell(i, 2).value = detail.get("reason")
                        ws_supp.cell(i, 3).value = detail.get("original")
                        ws_supp.cell(i, 4).value = detail.get("example")
                        ws_supp.cell(i, 5).value = detail.get("fixed")
                        ws_supp.cell(i, 6).value = detail.get("notes")
                    except (json.JSONDecodeError, AttributeError):
                        pass

            fname = f"PATCH_LIST_{month_str}_{version}.xlsx"
            p = str(out / fname)
            wb.save(p)
            paths.append(p)

        return paths

    def _find_test_report(self, base: Path, version: str, issue_no: str) -> str | None:
        """在 {base}/{version}/測試報告/ 搜尋含 issue_no 的 .doc/.docx，回傳絕對路徑。"""
        report_dir = base / version / "測試報告"
        if not report_dir.exists():
            return None
        for f in sorted(report_dir.iterdir()):
            if issue_no in f.name and f.suffix.lower() in (".doc", ".docx"):
                return str(f.absolute())
        return None

    def _write_patch_header_row(self, ws: object, headers: list[str], row: int = 1) -> None:
        """將標題列寫入指定行（可指定 row 參數，預設第 1 列）。"""
        from openpyxl.styles import Alignment, Font, PatternFill
        for c, h in enumerate(headers, start=1):
            cell = ws.cell(row, c)
            cell.value = h
            cell.font = Font(name="微軟正黑體", bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor=self._HDR_DARK)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def _parse_scan_meta(self, iss: PatchIssue) -> dict:
        """解析掃描模式的 mantis_detail，回傳 form_files/sql_files/muti_files。"""
        if not iss.mantis_detail:
            return {}
        try:
            return json.loads(iss.mantis_detail)
        except (json.JSONDecodeError, AttributeError):
            return {}

    def verify_patch_links(self, patch_dir: str) -> dict[str, dict]:
        """掃描 patch_dir 下各版本 PATCH_LIST_*.xlsx，驗證 Issue No 超連結是否有效。
        回傳 {version: {"total": N, "ok": N, "failed": [issue_no_or_path, ...]}}
        """
        import openpyxl

        base = Path(patch_dir)
        result: dict[str, dict] = {}
        for version_dir in sorted(base.iterdir()):
            if not version_dir.is_dir():
                continue
            version = version_dir.name
            xlsx_files = list(version_dir.glob("PATCH_LIST_*.xlsx"))
            if not xlsx_files:
                continue
            total = ok = 0
            failed: list[str] = []
            seen: set[str] = set()
            for xlsx_path in xlsx_files:
                try:
                    wb = openpyxl.load_workbook(str(xlsx_path))
                except Exception:
                    continue
                for sheet_name in ["IT 發行通知", "HR 發行通知"]:
                    if sheet_name not in wb.sheetnames:
                        continue
                    ws = wb[sheet_name]
                    for row in ws.iter_rows(min_row=2):
                        cell = row[0]
                        if not cell.value or cell.value in seen:
                            continue
                        seen.add(str(cell.value))
                        hl = cell.hyperlink
                        if hl is None:
                            continue
                        target = hl.target if hasattr(hl, "target") else str(hl)
                        total += 1
                        if target:
                            local = target.replace("file:///", "").replace("/", "\\")
                            if Path(local).exists():
                                ok += 1
                                continue
                        failed.append(f"{cell.value} → {target}")
            if total > 0:
                result[version] = {"total": total, "ok": ok, "failed": failed}
        return result

    def generate_patch_list_from_dir(
        self, patch_ids: dict[str, int], patch_dir: str, month_str: str
    ) -> list[str]:
        """依 patch_ids 各版本產 PATCH_LIST_{month_str}_{VER}.xlsx（IT/HR/補充說明格式），存至版本子目錄。"""
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill

        base = Path(patch_dir)
        paths: list[str] = []
        data_font = Font(name="微軟正黑體", size=11)
        data_align = Alignment(vertical="center", wrap_text=True)

        def _set_data_cell(cell, value):
            cell.value = value
            cell.font = data_font
            cell.alignment = data_align

        def _set_hyperlink_cell(cell, value, target):
            _set_data_cell(cell, value)
            cell.hyperlink = target
            cell.font = Font(name="微軟正黑體", size=11, color="0563C1", underline="single")

        for version, patch_id in patch_ids.items():
            issues = self._repo.list_issues_by_patch(patch_id)
            report_cache: dict[str, str | None] = {
                iss.issue_no: self._find_test_report(base, version, iss.issue_no)
                for iss in issues
            }
            wb = Workbook()

            # ① IT 發行通知
            ws_it = wb.active
            ws_it.title = "IT 發行通知"
            self._write_patch_header(
                ws_it, ["Issue No", "類型", "程式代號", "說明", "FORM 目錄", "DB 物件", "多語更新", "備註"]
            )
            for i, iss in enumerate(issues, start=2):
                meta = self._parse_scan_meta(iss)
                report_path = report_cache[iss.issue_no]
                clr = self._CLR_ENH if iss.issue_type == "Enhancement" else self._CLR_BUG
                row_fill = PatternFill("solid", fgColor=clr)
                if report_path:
                    _set_hyperlink_cell(ws_it.cell(i, 1), iss.issue_no, "file:///" + report_path.replace("\\", "/"))
                else:
                    _set_data_cell(ws_it.cell(i, 1), iss.issue_no)
                _set_data_cell(ws_it.cell(i, 2), iss.issue_type)
                _set_data_cell(ws_it.cell(i, 3), iss.program_code)
                _set_data_cell(ws_it.cell(i, 4), iss.description)
                _set_data_cell(ws_it.cell(i, 5), "、".join(meta.get("form_files") or []))
                _set_data_cell(ws_it.cell(i, 6), "、".join(meta.get("sql_files") or []))
                _set_data_cell(ws_it.cell(i, 7), "\n".join(meta.get("muti_files") or []))
                _set_data_cell(ws_it.cell(i, 8), None)
                for c in range(1, 9):
                    ws_it.cell(i, c).fill = row_fill

            # ② HR 發行通知
            ws_hr = wb.create_sheet("HR 發行通知")
            self._write_patch_header(ws_hr, [
                "Issue No", "計區域", "類型", "程式代號", "程式名稱",
                "功能說明", "影響說明/用途", "相關程式(FORM)",
                "上線所需動作", "測試方向及注意事項", "備註",
            ])
            for i, iss in enumerate(issues, start=2):
                meta = self._parse_scan_meta(iss)
                region_fill = PatternFill("solid", fgColor={
                    "TW": self._CLR_TW, "CN": self._CLR_CN
                }.get(iss.region or "", self._CLR_BOTH))
                report_path = report_cache[iss.issue_no]
                if report_path:
                    _set_hyperlink_cell(ws_hr.cell(i, 1), iss.issue_no, "file:///" + report_path.replace("\\", "/"))
                else:
                    _set_data_cell(ws_hr.cell(i, 1), iss.issue_no)
                _set_data_cell(ws_hr.cell(i, 2), iss.region)
                _set_data_cell(ws_hr.cell(i, 3), iss.issue_type)
                _set_data_cell(ws_hr.cell(i, 4), iss.program_code)
                _set_data_cell(ws_hr.cell(i, 5), iss.program_name)
                _set_data_cell(ws_hr.cell(i, 6), iss.description)
                _set_data_cell(ws_hr.cell(i, 7), iss.impact)
                _set_data_cell(ws_hr.cell(i, 8), "、".join(meta.get("form_files") or []))
                _set_data_cell(ws_hr.cell(i, 9), "請與資訊單位確認是否已完成更新\n確認更新完成再進行測試")
                ws_hr.cell(i, 9).fill = PatternFill("solid", fgColor="FFF9C4")
                _set_data_cell(ws_hr.cell(i, 10), iss.test_direction)
                _set_data_cell(ws_hr.cell(i, 11), None)
                for c in [1, 2, 3, 4, 5, 6, 7, 8, 10, 11]:
                    ws_hr.cell(i, c).fill = region_fill

            # ③ 問題修正補充說明（7 欄，含 Mantis supplement）
            ws_supp = wb.create_sheet("問題修正補充說明")
            self._write_patch_header(ws_supp, [
                "Issue No", "測試報告", "修改原因", "原問題", "範例說明", "修正後", "注意事項"
            ])
            for i, iss in enumerate(issues, start=2):
                meta = self._parse_scan_meta(iss)
                supplement = meta.get("supplement") or {}
                _set_data_cell(ws_supp.cell(i, 1), iss.issue_no)
                report_path = report_cache[iss.issue_no]
                cell2 = ws_supp.cell(i, 2)
                if report_path:
                    _set_hyperlink_cell(cell2, iss.issue_no, "file:///" + report_path.replace("\\", "/"))
                else:
                    _set_data_cell(cell2, iss.issue_no)
                _set_data_cell(ws_supp.cell(i, 3), supplement.get("修改原因", ""))
                _set_data_cell(ws_supp.cell(i, 4), supplement.get("原問題", ""))
                _set_data_cell(ws_supp.cell(i, 5), supplement.get("範例說明", ""))
                _set_data_cell(ws_supp.cell(i, 6), supplement.get("修正後", ""))
                _set_data_cell(ws_supp.cell(i, 7), supplement.get("注意事項", ""))

            fname = f"PATCH_LIST_{month_str}_{version}.xlsx"
            out_path = base / version / fname
            out_path.parent.mkdir(parents=True, exist_ok=True)
            try:
                wb.save(str(out_path))
            except PermissionError:
                raise PermissionError(f"無法寫入 {out_path.name}，請先關閉已開啟的 Excel 檔案後重試")
            paths.append(str(out_path))

        return paths

    def _write_patch_header(self, ws: object, headers: list[str]) -> None:
        self._write_patch_header_row(ws, headers, row=1)

    # ── 客戶通知 HTML ────────────────────────────────────────────────────────

    def generate_notify_html(
        self,
        patch_id: int,
        output_dir: str,
        month_str: str | None = None,
        notify_body: str | None = None,
        version: str = "11G",
        banner_image_bytes: bytes | None = None,
        schedule_reminders: list[str] | None = None,
    ) -> str:
        """使用 Jinja2 範本產客戶通知信 HTML。"""
        import base64
        from pathlib import Path as _Path

        from jinja2 import Environment, FileSystemLoader

        issues = self._repo.list_issues_by_patch(patch_id)
        if month_str is None:
            patch = self._repo.get_patch_by_id(patch_id)
            month_str = patch.month_str if patch and patch.month_str else "000000"

        year = month_str[:4]
        month_num = int(month_str[4:]) if month_str[4:].isdigit() else 1
        month = month_str[4:]
        color_dark, color_mid, season_icon = _SEASON_COLORS.get(month_num, ("1F4E79", "2E75B6", "🌸"))

        banner_b64 = None
        if banner_image_bytes:
            mime = "image/png" if banner_image_bytes[:4] == b"\x89PNG" else "image/jpeg"
            banner_b64 = f"data:{mime};base64,{base64.b64encode(banner_image_bytes).decode()}"

        templates_dir = _Path(__file__).parent.parent.parent.parent / "templates"
        env = Environment(loader=FileSystemLoader(str(templates_dir)), autoescape=True)
        tmpl = env.get_template("patch_notify.html.j2")

        html = tmpl.render(
            year=year,
            month=month,
            version=version,
            issues=issues,
            notify_body=notify_body or "",
            schedule_reminders=schedule_reminders or [],
            color_dark=color_dark,
            color_mid=color_mid,
            season_icon=season_icon,
            banner_b64=banner_b64,
        )

        out = _Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)
        fname = f"【HCP{version}維護客戶】{month_str}月份大PATCH更新通知.html"
        path = str(out / fname)
        _Path(path).write_text(html, encoding="utf-8")
        return path

    def generate_notify_html_from_dir(
        self,
        patch_ids: dict[str, int],
        patch_dir: str,
        month_str: str,
        schedule_reminders: list[str] | None = None,
        banner_image_bytes: bytes | None = None,
    ) -> list[str]:
        """依 patch_ids 各版本產客戶通知信 HTML，存至 patch_dir 頂層。"""
        paths: list[str] = []
        for version, patch_id in patch_ids.items():
            path = self.generate_notify_html(
                patch_id=patch_id,
                output_dir=patch_dir,
                month_str=month_str,
                version=version,
                schedule_reminders=schedule_reminders,
                banner_image_bytes=banner_image_bytes,
            )
            paths.append(path)
        return paths
