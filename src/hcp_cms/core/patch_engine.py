"""SinglePatchEngine — 單次 Patch 整理流程。"""

from __future__ import annotations

import re
import sqlite3
from pathlib import Path
from typing import TypedDict

from hcp_cms.data.models import PatchIssue, PatchRecord
from hcp_cms.data.repositories import PatchRepository

_ISSUE_TYPE_BUGFIX = "BugFix"
_ISSUE_TYPE_ENH = "Enhancement"


class PatchScanResult(TypedDict):
    """scan_patch_dir 回傳的資料夾掃描摘要。"""
    form_files: list[str]
    sql_files: list[str]
    muti_files: list[str]   # "muti" 是 HCP 系統內的真實子目錄名稱（多語更新）
    setup_bat: bool
    release_note: str | None
    install_guide: str | None
    missing: list[str]


class SinglePatchEngine:
    def __init__(self, conn: sqlite3.Connection) -> None:
        self._repo = PatchRepository(conn)

    # ── 掃描資料夾 ──────────────────────────────────────────────────────────

    def scan_patch_dir(self, patch_dir: str) -> PatchScanResult:
        """掃描解壓縮資料夾，回傳結構摘要。"""
        path = Path(patch_dir)
        result: PatchScanResult = {
            "form_files": [],
            "sql_files": [],
            "muti_files": [],
            "setup_bat": False,
            "release_note": None,
            "install_guide": None,
            "missing": [],
        }

        for sub, key, required in [("form", "form_files", True),
                                    ("sql", "sql_files", True),
                                    ("muti", "muti_files", False)]:
            sub_path = path / sub
            if sub_path.exists():
                result[key] = [f.name for f in sub_path.iterdir() if f.is_file()]
            elif required:
                result["missing"].append(f"{sub}/")

        result["setup_bat"] = (path / "setup.bat").exists()

        for f in path.iterdir():
            if not f.is_file():
                continue
            name_lower = f.name.lower()
            if "releasenote" in name_lower or "release_note" in name_lower:
                result["release_note"] = str(f)
            elif "installation" in name_lower or "installguide" in name_lower or "install_guide" in name_lower:
                result["install_guide"] = str(f)

        return result

    # ── 讀取 ReleaseNote ─────────────────────────────────────────────────────

    def read_release_doc(self, doc_path: str) -> list[dict[str, str]]:
        """解析 ReleaseNote.doc/.docx，回傳 Issue 清單。"""
        path = Path(doc_path)
        if not path.exists():
            return []
        text = self._extract_doc_text(path)
        if text is None:
            return []
        return self._parse_release_note_text(text)

    def _extract_doc_text(self, path: Path) -> str | None:
        if path.suffix.lower() == ".docx":
            return self._read_docx(path)
        return self._read_doc_win32(path)

    def _read_docx(self, path: Path) -> str | None:
        try:
            import docx as python_docx
            doc = python_docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return None

    def _read_doc_win32(self, path: Path) -> str | None:
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                d = word.Documents.Open(str(path.absolute()))
                text: str = d.Range().Text
                d.Close(False)
                return text
            finally:
                word.Quit()
        except Exception:
            return None

    def _parse_release_note_text(self, text: str) -> list[dict[str, str]]:
        """從文字中辨識 Issue 編號、類型、說明。"""
        issues: list[dict[str, str]] = []
        current_type = _ISSUE_TYPE_BUGFIX
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            low = stripped.lower()
            if "bug fix" in low or "bug修正" in low or "缺陷修正" in low:
                current_type = _ISSUE_TYPE_BUGFIX
                continue
            if "enhancement" in low or "功能加強" in low or "功能增強" in low:
                current_type = _ISSUE_TYPE_ENH
                continue
            m = re.match(r"(\d{7})\s+(.*)", stripped)
            if m:
                issues.append({
                    "issue_no": m.group(1),
                    "issue_type": current_type,
                    "description": m.group(2).strip(),
                    "region": "共用",
                })
        return issues

    # ── Excel 報表 ───────────────────────────────────────────────────────────

    _CLR_CS    = "D5F5E3"   # 客服驗證欄
    _CLR_CUST  = "D6EAF8"   # 客戶測試欄
    _CLR_PATCH = "FEF9E7"   # 可納入大Patch
    _CLR_NOTE  = "F5EEF8"   # 備註
    _CLR_ENH   = "E2EFDA"   # Enhancement 列
    _CLR_BUG   = "FCE4D6"   # BugFix 列
    _CLR_WARN  = "FFF3CD"   # 待確認

    def generate_excel_reports(self, patch_id: int, output_dir: str) -> list[str]:
        """產生 3 份 Excel：Issue清單整理、發行通知、IT_HR清單。"""
        from openpyxl import Workbook
        from openpyxl.styles import PatternFill

        issues = self._repo.list_issues_by_patch(patch_id)
        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)
        paths = []

        # ① Issue清單整理（內部追蹤）
        wb1 = Workbook()
        ws1 = wb1.active
        ws1.title = "Issue清單整理"
        hdrs1 = ["Issue No", "類型", "說明", "FORM", "SQL", "MUTI", "腳本",
                 "客服驗證", "客服測試結果", "客服測試日期",
                 "提供客戶驗證", "客戶測試結果", "客戶測試日期",
                 "可納入大Patch", "備註"]
        self._write_header_row(ws1, hdrs1)
        for i, iss in enumerate(issues, start=2):
            ws1.cell(i, 1).value = iss.issue_no
            ws1.cell(i, 2).value = iss.issue_type
            ws1.cell(i, 3).value = iss.description
            row_fill = self._CLR_ENH if iss.issue_type == "Enhancement" else self._CLR_BUG
            for c in range(1, 4):
                ws1.cell(i, c).fill = PatternFill("solid", fgColor=row_fill)
            for c in range(8, 11):
                ws1.cell(i, c).fill = PatternFill("solid", fgColor=self._CLR_CS)
            for c in range(11, 14):
                ws1.cell(i, c).fill = PatternFill("solid", fgColor=self._CLR_CUST)
            ws1.cell(i, 14).fill = PatternFill("solid", fgColor=self._CLR_PATCH)
            ws1.cell(i, 15).fill = PatternFill("solid", fgColor=self._CLR_NOTE)
        p1 = str(out / "Issue清單整理.xlsx")
        wb1.save(p1)
        paths.append(p1)

        # ② 發行通知（對客戶，不含追蹤欄）
        wb2 = Workbook()
        ws2 = wb2.active
        ws2.title = "發行通知"
        hdrs2 = ["Issue No", "類型", "說明", "FORM目錄", "DB物件", "多語更新", "安裝步驟", "備註"]
        self._write_header_row(ws2, hdrs2)
        for i, iss in enumerate(issues, start=2):
            ws2.cell(i, 1).value = iss.issue_no
            ws2.cell(i, 2).value = iss.issue_type
            ws2.cell(i, 3).value = iss.description
        p2 = str(out / "發行通知.xlsx")
        wb2.save(p2)
        paths.append(p2)

        # ③ IT/HR 清單（雙頁籤）
        wb3 = Workbook()
        ws_it = wb3.active
        ws_it.title = "IT 清單"
        hdrs_it = ["Issue No", "類型", "程式代號", "說明", "FORM目錄", "DB物件", "多語更新", "備註"]
        self._write_header_row(ws_it, hdrs_it)
        ws_hr = wb3.create_sheet("HR 清單")
        hdrs_hr = ["Issue No", "計區域", "類型", "程式代號", "程式名稱",
                   "功能說明", "影響說明", "相關程式(FORM)",
                   "上線所需動作", "測試方向及注意事項", "備註"]
        self._write_header_row(ws_hr, hdrs_hr)
        for i, iss in enumerate(issues, start=2):
            ws_it.cell(i, 1).value = iss.issue_no
            ws_it.cell(i, 2).value = iss.issue_type
            ws_it.cell(i, 3).value = iss.program_code
            ws_it.cell(i, 4).value = iss.description
            ws_hr.cell(i, 1).value = iss.issue_no
            ws_hr.cell(i, 2).value = iss.region
            ws_hr.cell(i, 3).value = iss.issue_type
            ws_hr.cell(i, 4).value = iss.program_code
            ws_hr.cell(i, 5).value = iss.program_name
            ws_hr.cell(i, 6).value = iss.description
            ws_hr.cell(i, 7).value = iss.impact
            ws_hr.cell(i, 9).value = "請與資訊單位確認是否已完成更新，確認更新完成再進行測試"
            ws_hr.cell(i, 10).value = iss.test_direction
        p3 = str(out / "Issue清單_IT_HR.xlsx")
        wb3.save(p3)
        paths.append(p3)

        return paths

    def _write_header_row(self, ws: object, headers: list[str]) -> None:
        from openpyxl.styles import Alignment, Font, PatternFill
        for c, h in enumerate(headers, start=1):
            cell = ws.cell(1, c)
            cell.value = h
            cell.font = Font(name="微軟正黑體", bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F3864")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # ── 測試腳本 ─────────────────────────────────────────────────────────────

    # ── 封存解壓縮 ────────────────────────────────────────────────────────────

    def load_from_archive(self, archive_path: str, output_dir: str) -> tuple[int, str, int]:
        """解壓 .7z 至 output_dir，掃描目錄，建立 DB Patch 記錄，讀 ReleaseNote。
        Returns (patch_id, version_tag, issue_count)
        """
        import py7zr
        archive = Path(archive_path)
        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)
        with py7zr.SevenZipFile(str(archive), mode="r") as z:
            z.extractall(path=str(out))
        version_tag = self._parse_version_tag(archive.name)
        patch = PatchRecord(type="single", patch_dir=str(out))
        patch_id = self._repo.insert_patch(patch)
        scan = self.scan_patch_dir(str(out))
        issue_count = 0
        if scan.get("release_note"):
            issue_count = self.load_issues_from_release_doc(patch_id, scan["release_note"])
        return patch_id, version_tag, issue_count

    def _parse_version_tag(self, filename: str) -> str:
        """從檔名解析版本標籤，優先比對 IP_合併_YYYYMMDD 格式。"""
        m = re.search(r"IP_合併_\d{8}", filename)
        if m:
            return m.group(0)
        stem = Path(filename).stem
        return stem[:20] if len(stem) > 20 else stem

    def extract_patch_archives(self, patch_dir: str) -> list[str]:
        """解壓縮資料夾內所有 .7z 封存檔至同一資料夾，回傳已解壓縮的檔名清單。"""
        import py7zr
        path = Path(patch_dir)
        extracted = []
        for archive in sorted(path.glob("*.7z")):
            try:
                with py7zr.SevenZipFile(str(archive), mode="r") as z:
                    z.extractall(path=str(path))
                extracted.append(archive.name)
            except Exception as e:
                import logging
                logging.warning("解壓縮失敗 [%s]: %s", archive.name, e)
        return extracted

    # ── Patch 記錄管理 ─────────────────────────────────────────────────────────

    def get_issue_nos_by_patch(self, patch_id: int) -> list[str]:
        """取得 Patch 所有非空 Issue 編號。"""
        issues = self._repo.list_issues_by_patch(patch_id)
        return [i.issue_no for i in issues if i.issue_no]

    def setup_new_patch(self, patch_dir: str) -> int:
        """建立單次 Patch 記錄，回傳 patch_id。"""
        patch = PatchRecord(type="single", patch_dir=patch_dir)
        return self._repo.insert_patch(patch)

    def load_issues_from_release_doc(self, patch_id: int, doc_path: str) -> int:
        """從 ReleaseNote 解析 Issues 並寫入 DB，回傳新增筆數。"""
        raw_list = self.read_release_doc(doc_path)
        for idx, raw in enumerate(raw_list):
            issue = PatchIssue(
                patch_id=patch_id,
                issue_no=raw.get("issue_no", ""),
                program_code=raw.get("program_code"),
                program_name=raw.get("program_name"),
                issue_type=raw.get("issue_type", "BugFix"),
                region=raw.get("region", "共用"),
                description=raw.get("description"),
                impact=raw.get("impact"),
                test_direction=raw.get("test_direction"),
                source="manual",
                sort_order=idx,
            )
            self._repo.insert_issue(issue)
        return len(raw_list)

    # ── 測試腳本 ─────────────────────────────────────────────────────────────

    def generate_test_scripts(self, patch_id: int, output_dir: str) -> list[str]:
        """產生測試腳本_客服版.docx、客戶版.docx、測試追蹤表.xlsx。"""
        import docx as python_docx
        from openpyxl import Workbook

        issues = self._repo.list_issues_by_patch(patch_id)
        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)
        paths = []

        # 客服版 Word
        doc_cs = python_docx.Document()
        doc_cs.add_heading("測試腳本（客服版）", 0)
        for iss in issues:
            doc_cs.add_heading(f"Issue {iss.issue_no}", level=1)
            doc_cs.add_paragraph(f"說明：{iss.description or ''}")
            doc_cs.add_paragraph(f"測試步驟：{iss.test_direction or '請填寫'}")
            doc_cs.add_paragraph("測試人員：＿＿＿＿　審核人員：＿＿＿＿")
        p_cs = str(out / "測試腳本_客服版.docx")
        doc_cs.save(p_cs)
        paths.append(p_cs)

        # 客戶版 Word
        doc_cu = python_docx.Document()
        doc_cu.add_heading("測試腳本（客戶版）", 0)
        for iss in issues:
            doc_cu.add_heading(f"Issue {iss.issue_no}", level=1)
            doc_cu.add_paragraph(f"說明：{iss.description or ''}")
            doc_cu.add_paragraph("□ 正常　□ 異常")
            doc_cu.add_paragraph("客戶回覆日期：＿＿＿＿　簽名：＿＿＿＿")
        p_cu = str(out / "測試腳本_客戶版.docx")
        doc_cu.save(p_cu)
        paths.append(p_cu)

        # 追蹤表 xlsx
        wb = Workbook()
        ws_cs = wb.active
        ws_cs.title = "客服驗證"
        self._write_header_row(ws_cs, ["Issue No", "說明", "測試結果(PASS/FAIL)", "測試日期", "備註"])
        ws_cu = wb.create_sheet("客戶驗證")
        self._write_header_row(ws_cu, ["Issue No", "說明", "測試結果(正常/異常)", "回覆日期", "備註"])
        for i, iss in enumerate(issues, start=2):
            ws_cs.cell(i, 1).value = iss.issue_no
            ws_cs.cell(i, 2).value = iss.description
            ws_cu.cell(i, 1).value = iss.issue_no
            ws_cu.cell(i, 2).value = iss.description
        p_tr = str(out / "測試追蹤表.xlsx")
        wb.save(p_tr)
        paths.append(p_tr)

        return paths
