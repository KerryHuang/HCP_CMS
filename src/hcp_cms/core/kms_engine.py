"""KMS knowledge management — search, CRUD, auto-extract, Excel import/export."""

import sqlite3
from pathlib import Path

import openpyxl

from hcp_cms.core.anonymizer import Anonymizer
from hcp_cms.data.fts import FTSManager
from hcp_cms.data.models import Case, QAKnowledge
from hcp_cms.data.repositories import QARepository
from hcp_cms.services.mail.base import RawEmail

# Question detection patterns
_QUESTION_PATTERNS = ["請問", "如何", "是否", "怎麼", "可否", "能否", "為什麼", "為何"]

# 圖片副檔名（與 msg_reader._IMAGE_EXTS 保持一致）
_IMAGE_EXTS = frozenset({".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"})


class KMSEngine:
    def __init__(self, conn: sqlite3.Connection) -> None:
        self._conn = conn
        self._qa_repo = QARepository(conn)
        self._fts = FTSManager(conn)
        self._anonymizer = Anonymizer()

    def search(self, query: str, system_product: str | None = None) -> list[QAKnowledge]:
        """Search QA via FTS5, return full QAKnowledge objects. 只返回已完成的 QA。"""
        fts_results = self._fts.search_qa(query)
        qa_list = []
        for result in fts_results:
            qa = self._qa_repo.get_by_id(result["qa_id"])
            if qa:
                if qa.status != "已完成":
                    continue
                if system_product and qa.system_product != system_product:
                    continue
                qa_list.append(qa)
        return qa_list

    def create_qa(
        self,
        question: str,
        answer: str,
        solution: str | None = None,
        keywords: str | None = None,
        system_product: str | None = None,
        issue_type: str | None = None,
        error_type: str | None = None,
        source: str = "manual",
        source_case_id: str | None = None,
        created_by: str | None = None,
        status: str = "已完成",
    ) -> QAKnowledge:
        """Create a new QA entry. 若 status 為已完成則建立 FTS 索引。"""
        qa_id = self._qa_repo.next_qa_id()
        qa = QAKnowledge(
            qa_id=qa_id,
            question=question,
            answer=answer,
            solution=solution,
            keywords=keywords,
            system_product=system_product,
            issue_type=issue_type,
            error_type=error_type,
            source=source,
            source_case_id=source_case_id,
            created_by=created_by,
            status=status,
        )
        self._qa_repo.insert(qa)
        if status == "已完成":
            self._fts.index_qa(qa_id, question, answer, solution, keywords)
        return qa

    def update_qa(self, qa_id: str, **fields: str | None) -> QAKnowledge | None:
        """Update QA fields and rebuild FTS index. 已完成 → 待審核 降級被拒絕。"""
        qa = self._qa_repo.get_by_id(qa_id)
        if not qa:
            return None
        incoming_status = fields.get("status")
        if qa.status == "已完成" and incoming_status == "待審核":
            return None
        for key, value in fields.items():
            if hasattr(qa, key):
                setattr(qa, key, value)
        self._qa_repo.update(qa)
        if qa.status == "已完成":
            self._fts.update_qa_index(qa_id, qa.question, qa.answer, qa.solution, qa.keywords)
        return qa

    def delete_qa(self, qa_id: str) -> None:
        """Delete QA and remove FTS index."""
        self._qa_repo.delete(qa_id)
        self._fts.remove_qa_index(qa_id)

    def extract_qa_from_email(
        self,
        raw_email: RawEmail,
        case_id: str | None = None,
        db_dir: Path | None = None,
    ) -> QAKnowledge | None:
        """從 RawEmail thread 欄位抽取 QA，儲存為待審核。無問題段則回傳 None。"""
        if not raw_email.thread_question:
            return None
        qa = self.create_qa(
            question=raw_email.thread_question,
            answer=raw_email.thread_answer or "",
            source="email",
            source_case_id=case_id,
            status="待審核",
        )
        if db_dir is not None and raw_email.source_file:
            self.attach_images(qa.qa_id, Path(raw_email.source_file), db_dir)
            qa = self._qa_repo.get_by_id(qa.qa_id) or qa
        return qa

    def approve_qa(self, qa_id: str, **updated_fields) -> QAKnowledge | None:
        """單一入口：更新欄位 → status='已完成' → FTS 索引建立。"""
        updated_fields["status"] = "已完成"
        return self.update_qa(qa_id, **updated_fields)

    def list_pending(self) -> list[QAKnowledge]:
        """列出所有待審核 QA。"""
        return self._qa_repo.list_by_status("待審核")

    def list_approved(self) -> list[QAKnowledge]:
        """列出所有已完成 QA（供 UI 層顯示全部使用）。"""
        return self._qa_repo.list_approved()

    def auto_extract_qa(
        self,
        case: Case,
        company_domain: str = "",
        company_aliases: list[str] | None = None,
    ) -> QAKnowledge | None:
        """Auto-extract QA from a case if it contains a question pattern.

        .. deprecated:: 由 extract_qa_from_email() 取代。
        Returns QA with source='email' or None if no question detected."""
        if not case.subject:
            return None

        text = f"{case.subject} {case.progress or ''}"
        has_question = any(p in text for p in _QUESTION_PATTERNS)
        if not has_question:
            return None

        question = self._anonymizer.anonymize(case.subject, company_domain, company_aliases)
        answer = self._anonymizer.anonymize(case.progress or "", company_domain, company_aliases)

        return self.create_qa(
            question=question,
            answer=answer,
            system_product=case.system_product,
            issue_type=case.issue_type,
            error_type=case.error_type,
            source="email",
            source_case_id=case.case_id,
        )

    def import_from_excel(self, file_path: Path) -> int:
        """Import QA entries from Excel file. Returns count imported."""
        wb = openpyxl.load_workbook(str(file_path), read_only=True)
        ws = wb.active
        count = 0

        rows = list(ws.iter_rows(min_row=2, values_only=True))  # Skip header
        for row in rows:
            if len(row) < 2 or not row[0]:
                continue
            question = str(row[0]) if row[0] else ""
            answer = str(row[1]) if len(row) > 1 and row[1] else ""
            solution = str(row[2]) if len(row) > 2 and row[2] else None
            keywords = str(row[3]) if len(row) > 3 and row[3] else None

            if question:
                self.create_qa(question=question, answer=answer, solution=solution, keywords=keywords, source="import")
                count += 1

        wb.close()
        return count

    def attach_images(self, qa_id: str, msg_path: Path, db_dir: Path) -> int:
        """從 msg_path 提取圖片至 db_dir/kms_attachments/qa_id/。
        若 msg_path 不存在回傳 0 且不更新 DB。
        """
        if not msg_path.exists():
            return 0
        from hcp_cms.services.mail.msg_reader import MSGReader
        dest_dir = db_dir / "kms_attachments" / qa_id
        saved = MSGReader.extract_images(msg_path, dest_dir)
        # 無論有無圖片，只要 msg 存在就更新 DB（避免重複嘗試）
        qa = self._qa_repo.get_by_id(qa_id)
        if qa:
            qa.has_image = "是"
            qa.doc_name = str(msg_path)
            self._qa_repo.update(qa)
        return len(saved)

    def export_to_docx(
        self,
        file_path: Path,
        db_dir: Path,
        qa_list: list[QAKnowledge] | None = None,
    ) -> Path:
        """匯出 QA 至 Word 文件。qa_list=None 時匯出全部已完成 QA。"""
        from docx import Document
        from docx.shared import Cm

        if qa_list is None:
            qa_list = self._qa_repo.list_approved()

        doc = Document()
        doc.add_heading("KMS 知識庫匯出", level=1)

        if not qa_list:
            doc.add_paragraph("無資料")
            doc.save(str(file_path))
            return file_path

        for i, qa in enumerate(qa_list):
            title = f"{qa.qa_id}"
            if qa.system_product:
                title += f"｜{qa.system_product}"
            doc.add_heading(title, level=2)
            doc.add_paragraph(f"問題：{qa.question or ''}")
            doc.add_paragraph(f"回覆：{qa.answer or ''}")

            # 插入圖片
            img_dir = db_dir / "kms_attachments" / qa.qa_id
            if img_dir.exists():
                for img_path in sorted(img_dir.iterdir()):
                    if img_path.suffix.lower() in _IMAGE_EXTS:
                        try:
                            doc.add_picture(str(img_path), width=Cm(14))
                        except Exception:
                            continue

            if qa.solution:
                doc.add_paragraph(f"解決方案：{qa.solution}")
            if i < len(qa_list) - 1:
                doc.add_paragraph("─" * 40)

        doc.save(str(file_path))
        return file_path

    def export_to_excel(self, file_path: Path, qa_list: list[QAKnowledge] | None = None) -> Path:
        """Export QA entries to Excel. If qa_list is None, export approved only."""
        if qa_list is None:
            qa_list = self._qa_repo.list_approved()

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "QA知識庫"

        # Header
        headers = ["QA編號", "問題", "回覆", "解決方案", "關鍵字", "產品", "問題類型", "錯誤類型", "來源", "建立日期"]
        ws.append(headers)

        for qa in qa_list:
            ws.append([
                qa.qa_id, qa.question, qa.answer, qa.solution, qa.keywords,
                qa.system_product, qa.issue_type, qa.error_type, qa.source, qa.created_at,
            ])

        wb.save(str(file_path))
        return file_path
